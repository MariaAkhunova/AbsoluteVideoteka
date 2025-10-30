# views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.db.models import Count, Q
from django.contrib import messages
from django.contrib.auth import login, logout
from django.contrib.auth.models import User
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.template.loader import render_to_string
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
from docx import Document
from docx.shared import Inches
import os
from .models import Sale, MovieCrew, Movie, Artist, Role
from datetime import date
from .forms import UserEmailForm, PurchaseForm, MovieFilterForm

# Главная страница с фильтрацией и сортировкой
def homepage(request):
    movies_list = Movie.objects.all()
    
    # Применяем фильтрацию
    form = MovieFilterForm(request.GET)
    if form.is_valid():
        genre = form.cleaned_data.get('genre')
        min_year = form.cleaned_data.get('min_year')
        max_year = form.cleaned_data.get('max_year')
        min_price = form.cleaned_data.get('min_price')
        max_price = form.cleaned_data.get('max_price')
        search = form.cleaned_data.get('search')
        sort_by = form.cleaned_data.get('sort_by', 'title')
        
        # Поиск
        if search:
            movies_list = movies_list.filter(
                Q(title__icontains=search) | 
                Q(description__icontains=search)
            )
        
        # Фильтрация по жанру
        if genre:
            movies_list = movies_list.filter(genre__icontains=genre)
        
        # Фильтрация по году
        if min_year:
            movies_list = movies_list.filter(release_year__gte=min_year)
        if max_year:
            movies_list = movies_list.filter(release_year__lte=max_year)
        
        # Фильтрация по цене
        if min_price:
            movies_list = movies_list.filter(price__gte=min_price)
        if max_price:
            movies_list = movies_list.filter(price__lte=max_price)
        
        # Сортировка
        if sort_by == 'title':
            movies_list = movies_list.order_by('title')
        elif sort_by == 'year_desc':
            movies_list = movies_list.order_by('-release_year')
        elif sort_by == 'year_asc':
            movies_list = movies_list.order_by('release_year')
        elif sort_by == 'price_desc':
            movies_list = movies_list.order_by('-price')
        elif sort_by == 'price_asc':
            movies_list = movies_list.order_by('price')
    
    # Экспорт данных
    export_format = request.GET.get('export')
    if export_format and movies_list.exists():
        return export_movies(movies_list, export_format)
    
    context = {
        'movies_list': movies_list,
        'form': form,
        'results_count': movies_list.count()
    }
    return render(request, 'home.html', context)

# Функция экспорта с исправленной кодировкой
def export_movies(movies, format_type):
    
    if format_type == 'txt':
        # Указываем правильную кодировку для TXT
        response = HttpResponse(content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = 'attachment; filename="movies.txt"'
        
        content = "Список фильмов\n"
        content += "=" * 50 + "\n\n"
        
        for i, movie in enumerate(movies, 1):
            content += f"{i}. {movie.title}\n"
            content += f"   Год выпуска: {movie.release_year}\n"
            content += f"   Жанр: {movie.genre}\n"
            content += f"   Цена: {movie.price} руб.\n"
            if movie.duration:
                content += f"   Длительность: {movie.duration} мин.\n"
            content += "\n"
        
        # Кодируем в UTF-8 и записываем BOM для правильного отображения
        response.write(content.encode('utf-8-sig'))
        return response
    
    elif format_type == 'pdf':
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="movies.pdf"'
        
        buffer = io.BytesIO()
        
        try:
            # Пробуем зарегистрировать шрифт с поддержкой кириллицы
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
        except:
            pass
        
        # Создаем документ
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        elements = []
        
        # Стили для русского текста
        styles = getSampleStyleSheet()
        
        # Создаем кастомные стили с шрифтами, поддерживающими кириллицу
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=16,
            spaceAfter=30,
            alignment=1  # center
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            spaceAfter=12
        )
        
        # Заголовок
        title = Paragraph("Список фильмов", title_style)
        elements.append(title)
        
        # Информация о количестве
        info_text = Paragraph(f"Всего фильмов: {movies.count()}", normal_style)
        elements.append(info_text)
        elements.append(Spacer(1, 20))
        
        # Данные для таблицы
        data = [['Название', 'Год', 'Жанр', 'Цена', 'Длительность']]
        
        for movie in movies:
            duration = f"{movie.duration} мин." if movie.duration else "—"
            data.append([
                movie.title,
                str(movie.release_year),
                movie.genre,
                f"{movie.price} руб.",
                duration
            ])
        
        # Создание таблицы
        table = Table(data, colWidths=[200, 50, 80, 60, 70])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (1, -1), 'CENTER'),
            ('ALIGN', (3, 1), (4, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
        ]))
        
        elements.append(table)
        doc.build(elements)
        
        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)
        return response
    
    elif format_type == 'docx':
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="movies.docx"'
        
        # Создаем документ Word
        doc = Document()
        
        # Добавляем заголовок
        title = doc.add_heading('Список фильмов', 0)
        
        # Информация о количестве
        doc.add_paragraph(f'Всего фильмов: {movies.count()}')
        doc.add_paragraph()
        
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Год'
        hdr_cells[2].text = 'Жанр'
        hdr_cells[3].text = 'Цена'
        hdr_cells[4].text = 'Длительность'
        
        # Делаем заголовки жирными
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Добавляем данные
        for movie in movies:
            row_cells = table.add_row().cells
            row_cells[0].text = movie.title
            row_cells[1].text = str(movie.release_year)
            row_cells[2].text = movie.genre
            row_cells[3].text = f"{movie.price} руб."
            row_cells[4].text = f"{movie.duration} мин." if movie.duration else "—"
        
        doc.save(response)
        return response
    
    return redirect('home')

# Остальные функции остаются без изменений
@login_required
def purchase_movie(request, movie_id):
    movie = get_object_or_404(Movie, movie_id=movie_id)
    user = request.user
    
    if Sale.objects.filter(user=user, movie=movie).exists():
        messages.warning(request, f'Фильм "{movie.title}" уже куплен!')
        return redirect('movie_detail', movie_id=movie_id)
    
    if request.method == 'POST':
        form = PurchaseForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            
            if user.email != email:
                user.email = email
                user.save()
                messages.info(request, 'Email успешно обновлен!')
            
            sale = Sale.objects.create(
                user=user,
                movie=movie,
                unit_price=movie.price
            )
            
            messages.success(request, f'Фильм "{movie.title}" успешно куплен!')
            return redirect('profile')
    else:
        initial_email = user.email if user.email else ''
        form = PurchaseForm(initial={'email': initial_email})
    
    context = {
        'movie': movie,
        'form': form,
    }
    return render(request, 'purchase.html', context)

@login_required
def profile(request):
    user = request.user
    purchases = Sale.objects.filter(user=user).select_related('movie').order_by('-sale_date')
    
    if request.method == 'POST':
        form = UserEmailForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Email успешно обновлен!')
            return redirect('profile')
    else:
        form = UserEmailForm(instance=user)
    
    context = {
        'form': form,
        'purchases': purchases,
    }
    return render(request, 'profile.html', context)

def register_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('home')
    else:
        form = UserCreationForm()
    return render(request, 'register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('home')
    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})

def logout_view(request):
    logout(request)
    return redirect('home')    

def movie_detail(request, movie_id):
    movie = get_object_or_404(Movie, movie_id=movie_id)
    crew_by_roles = {}
    for crew in movie.moviecrew_set.all().select_related('artist', 'role'):
        role_name = crew.role.role_name
        if role_name not in crew_by_roles:
            crew_by_roles[role_name] = []
        crew_by_roles[role_name].append(crew)
    
    context = {
        'movie': movie,
        'crew_by_roles': crew_by_roles,
        'id': movie_id,
    }
    return render(request, 'movie_detail.html', context)  

def artist_detail(request, artist_id):
    artist = get_object_or_404(Artist.objects.prefetch_related(
        'moviecrew_set__movie',
        'moviecrew_set__role'
    ), artist_id=artist_id)
    
    movies_by_role = {}
    for crew in artist.moviecrew_set.all():
        role_name = crew.role.role_name
        if role_name not in movies_by_role:
            movies_by_role[role_name] = []
        movies_by_role[role_name].append({
            'movie': crew.movie,
            'character_name': crew.character_name
        })
    
    role_stats = artist.moviecrew_set.values(
        'role__role_name'
    ).annotate(
        count=Count('artist_id')
    ).order_by('-count')
    
    context = {
        'artist': artist,
        'movies_by_role': movies_by_role,
        'role_stats': role_stats,
        'age': artist.get_age(),
    }
    return render(request, 'artist_detail.html', context)
#покупка фильма
@login_required
def purchase_movie(request, movie_id):
    movie = get_object_or_404(Movie, movie_id=movie_id)
    user = request.user
    
    #проверка не куплен ли фильм
    if Sale.objects.filter(user=user, movie=movie).exists():
        messages.warning(request, f'Фильм "{movie.title}" уже куплен!')
        return redirect('movie_detail', movie_id=movie_id)
    
    if request.method == 'POST':
        form = PurchaseForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            
            #обновление почты пользователя если она изменилась
            if user.email != email:
                user.email = email
                user.save()
                messages.info(request, 'Email успешно обновлен!')
            
            #создание записи о покупке
            sale = Sale.objects.create(
                user=user,
                movie=movie,
                unit_price=movie.price
            )
            
            messages.success(request, f'Фильм "{movie.title}" успешно куплен!')
            return redirect('profile')
    else:
        #заполнение почты пользователя
        initial_email = user.email if user.email else ''
        form = PurchaseForm(initial={'email': initial_email})
    
    context = {
        'movie': movie,
        'form': form,
    }
    return render(request, 'purchase.html', context)

#профиль пользователя
@login_required
def profile(request):
    user = request.user
    purchases = Sale.objects.filter(user=user).select_related('movie').order_by('-sale_date')
    
    #обновление почты пользователя
    if request.method == 'POST':
        form = UserEmailForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Email успешно обновлен!')
            return redirect('profile')
    else:
        form = UserEmailForm(instance=user)
    
    context = {
        'form': form,
        'purchases': purchases,
    }
    return render(request, 'profile.html', context)

#регистрация пользоателя
def register_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('home')
    else:
        form = UserCreationForm()
    return render(request, 'register.html', {'form': form})

#вход в аккаунт пользователя
def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('home')
    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})

#выход из аккаунта пользователя
def logout_view(request):
    logout(request)
    return redirect('home')    

#показ подробностей о фильме
def movie_detail(request, movie_id):
    movie = get_object_or_404(Movie, movie_id=movie_id)
    crew_by_roles = {}
    for crew in movie.moviecrew_set.all().select_related('artist', 'role'):
        role_name = crew.role.role_name
        if role_name not in crew_by_roles:
            crew_by_roles[role_name] = []
        crew_by_roles[role_name].append(crew)
    
    context = {
        'movie': movie,
        'crew_by_roles': crew_by_roles,
        'id': movie_id,
    }
    return render(request, 'movie_detail.html', context)  

#подробности о артистах
def artist_detail(request, artist_id):
    artist = get_object_or_404(Artist.objects.prefetch_related(
        'moviecrew_set__movie',
        'moviecrew_set__role'
    ), artist_id=artist_id)
    
    #фильмы с группировкой по ролям
    movies_by_role = {}
    for crew in artist.moviecrew_set.all():
        role_name = crew.role.role_name
        if role_name not in movies_by_role:
            movies_by_role[role_name] = []
        movies_by_role[role_name].append({
            'movie': crew.movie,
            'character_name': crew.character_name
        })
    
    #статистика по ролям
    role_stats = artist.moviecrew_set.values(
        'role__role_name'
    ).annotate(
        count=Count('artist_id')
    ).order_by('-count')
    
    context = {
        'artist': artist,
        'movies_by_role': movies_by_role,
        'role_stats': role_stats,
        'age': artist.get_age(),
    }
    return render(request, 'artist_detail.html', context)   